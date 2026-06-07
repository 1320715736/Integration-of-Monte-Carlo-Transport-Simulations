from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


BASE_DIR = Path(__file__).resolve().parent
PDF_DIR = BASE_DIR.parent / "英文PDF"
SOURCE_MD = PDF_DIR / "paper_v2_english.md"
DOCX_INPUT_MD = BASE_DIR / "paper_v2_english_docx_input.md"
REFERENCE_DOCX = BASE_DIR / "mej_reference.docx"
PANDOC_DOCX = BASE_DIR / "paper_v2_english_MEJ_pandoc.docx"
CANDIDATE_DOCX = BASE_DIR / "paper_v2_english_MEJ_candidate.docx"
OUTPUT_DOCX = BASE_DIR / "paper_v2_english_MEJ.docx"


def set_run_font(run, name: str = "Times New Roman", size_pt: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def configure_paragraph_format(style, *, before=0, after=6, line=1.15, first_line=0.0, keep_next=False, align=None) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.first_line_indent = Inches(first_line) if first_line else None
    fmt.keep_with_next = keep_next
    if align is not None:
        fmt.alignment = align


def clear_paragraph_indent(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.left_indent = Inches(0)


def configure_style(style, *, size, bold=False, italic=False, align=None, before=0, after=6, line=1.15, first_line=0.0, keep_next=False) -> None:
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor(0, 0, 0)
    configure_paragraph_format(style, before=before, after=after, line=line, first_line=first_line, keep_next=keep_next, align=align)


def remove_style_border(style) -> None:
    p_pr = style._element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def remove_paragraph_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def build_reference_docx() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.20)
    sec.bottom_margin = Cm(2.20)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    sec.header_distance = Cm(1.25)
    sec.footer_distance = Cm(1.25)

    styles = doc.styles
    configure_style(styles["Normal"], size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=3.5, line=1.08, first_line=0.28)
    configure_style(styles["Title"], size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12, line=1.08)
    remove_style_border(styles["Title"])
    configure_style(styles["Heading 1"], size=14, bold=True, before=12, after=5, line=1.08, keep_next=True)
    configure_style(styles["Heading 2"], size=13, bold=True, before=10, after=4, line=1.08, keep_next=True)
    configure_style(styles["Heading 3"], size=12, bold=True, italic=False, before=7, after=3, line=1.08, keep_next=True)
    configure_style(styles["Caption"], size=10.5, before=2, after=5, line=1.03)
    configure_style(styles["Body Text"], size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=3.5, line=1.08, first_line=0.28)

    for style_name in ("Table Caption", "Image Caption"):
        if style_name in styles:
            configure_style(styles[style_name], size=10.5, before=4, after=6, line=1.05)

    doc.add_paragraph("Reference", style="Title")
    doc.save(REFERENCE_DOCX)


SUPERSCRIPT = str.maketrans("0123456789+-=()", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾")
SUBSCRIPT = str.maketrans({
    "0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄",
    "5": "₅", "6": "₆", "7": "₇", "8": "₈", "9": "₉",
    "+": "₊", "-": "₋", "=": "₌", "(": "₍", ")": "₎",
    "a": "ₐ", "e": "ₑ", "h": "ₕ", "i": "ᵢ", "j": "ⱼ",
    "k": "ₖ", "l": "ₗ", "m": "ₘ", "n": "ₙ", "o": "ₒ",
    "p": "ₚ", "r": "ᵣ", "s": "ₛ", "t": "ₜ", "u": "ᵤ",
    "v": "ᵥ", "x": "ₓ",
})


def super_text(value: str) -> str:
    return value.translate(SUPERSCRIPT)


def sub_text(value: str) -> str:
    if "/" in value:
        return value
    return value.translate(SUBSCRIPT)


def inline_math_to_text(match: re.Match) -> str:
    expr = match.group(1)
    text = expr
    text = text.replace(r"\mathrm{", "{")
    text = text.replace(r"\text{", "{")
    text = text.replace(r"\mu", "μ")
    text = text.replace(r"\times", "×")
    text = text.replace(r"\approx", "≈")
    text = text.replace(r"\leq", "≤").replace(r"\geq", "≥")
    text = text.replace(r"\varepsilon", "ε")
    text = text.replace(r"\sigma", "σ")
    text = text.replace(r"\tau", "τ")
    text = text.replace(r"\ ", " ")

    # Common detector notation reads better as editable text in Word.
    text = text.replace("p^+", "p+").replace("n^+", "n+")

    def replace_braced_sup(m: re.Match) -> str:
        return super_text(m.group(1))

    def replace_simple_sup(m: re.Match) -> str:
        return super_text(m.group(1))

    def replace_braced_sub(m: re.Match) -> str:
        return sub_text(m.group(1).replace("{", "").replace("}", ""))

    def replace_simple_sub(m: re.Match) -> str:
        return sub_text(m.group(1))

    text = re.sub(r"\^\{([^{}]+)\}", replace_braced_sup, text)
    text = re.sub(r"\^([0-9+\-=()])", replace_simple_sup, text)
    text = re.sub(r"_\{([^{}]+)\}", replace_braced_sub, text)
    text = re.sub(r"_([A-Za-z0-9])", replace_simple_sub, text)
    text = text.replace("{", "").replace("}", "")
    text = text.replace("\\", "")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def prepare_docx_markdown() -> None:
    text = SOURCE_MD.read_text(encoding="utf-8")

    # Convert only inline math spans. Display equations remain as LaTeX so
    # Pandoc can keep them as editable Word equations.
    text = re.sub(r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)", inline_math_to_text, text)
    DOCX_INPUT_MD.write_text(text, encoding="utf-8")


def run_pandoc() -> None:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("pandoc not found in PATH")
    cmd = [
        pandoc,
        str(DOCX_INPUT_MD),
        "-f",
        "markdown+tex_math_dollars+pipe_tables+raw_tex",
        "-t",
        "docx",
        "--reference-doc",
        str(REFERENCE_DOCX),
        "-o",
        str(PANDOC_DOCX),
    ]
    subprocess.run(cmd, cwd=PDF_DIR, check=True)


def set_tbl_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    settings = {
        "top": ("single", "000000", "8"),
        "bottom": ("single", "000000", "8"),
        "insideH": ("single", "BFBFBF", "4"),
        "left": ("nil", "auto", "0"),
        "right": ("nil", "auto", "0"),
        "insideV": ("nil", "auto", "0"),
    }
    for edge, (val, color, size) in settings.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), val)
        element.set(qn("w:color"), color)
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")


def set_table_row_rules(table) -> None:
    # Published MEJ tables use the standard Elsevier look: horizontal rules
    # only, no vertical rules and no shaded header.
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)

            for edge in ("left", "right"):
                element = borders.find(qn(f"w:{edge}"))
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    borders.append(element)
                element.set(qn("w:val"), "nil")
                element.set(qn("w:color"), "auto")
                element.set(qn("w:sz"), "0")
                element.set(qn("w:space"), "0")

            for edge in ("top", "bottom"):
                element = borders.find(qn(f"w:{edge}"))
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    borders.append(element)

                if row_i == 0 and edge == "top":
                    val, color, size = "single", "000000", "6"
                elif row_i == 0 and edge == "bottom":
                    val, color, size = "single", "000000", "4"
                elif row_i == len(table.rows) - 1 and edge == "bottom":
                    val, color, size = "single", "000000", "6"
                elif edge == "bottom":
                    val, color, size = "single", "BFBFBF", "2"
                else:
                    val, color, size = "nil", "auto", "0"

                element.set(qn("w:val"), val)
                element.set(qn("w:color"), color)
                element.set(qn("w:sz"), size)
                element.set(qn("w:space"), "0")


def clear_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")
        element.set(qn("w:color"), "auto")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ("top", "bottom", "left", "right"):
                element = tc_borders.find(qn(f"w:{edge}"))
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    tc_borders.append(element)
                element.set(qn("w:val"), "nil")
                element.set(qn("w:color"), "auto")
                element.set(qn("w:sz"), "0")
                element.set(qn("w:space"), "0")


def set_cell_margins(table, top=70, start=120, bottom=70, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        child = margins.find(qn(f"w:{tag}"))
        if child is None:
            child = OxmlElement(f"w:{tag}")
            margins.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in: list[float]) -> None:
    widths = [int(w * 1440) for w in widths_in]
    total = sum(widths)
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def clear_cell_shading(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is not None:
        tc_pr.remove(shd)


TABLE_DATA = [
    [
        ["Parameter", "Value"],
        ["p+ thickness", "0.2 μm"],
        ["i-region thickness", "5, 8, 10-130, 150, and 180 μm"],
        ["n+ thickness", "0.5 μm"],
        ["Lateral size", "240 μm × 240 μm"],
        ["Net i-region doping", "5.6 × 10¹² cm⁻³"],
        ["p+ / n+ doping", "1 × 10¹⁹ cm⁻³"],
    ],
    [
        ["Source", "Neh(Geant4)", "Neh(TCAD)", "Relative error (%)"],
        ["20 keV", "2.353 × 10³", "2.353 × 10³", "-1.93 × 10⁻¹⁴"],
        ["49 keV", "5.805 × 10³", "5.805 × 10³", "3.13 × 10⁻¹⁴"],
        ["100 keV", "1.191 × 10⁴", "1.191 × 10⁴", "0"],
        ["156.5 keV", "1.875 × 10⁴", "1.875 × 10⁴", "-1.94 × 10⁻¹⁴"],
        ["¹⁴C spectrum", "5.919 × 10³", "5.919 × 10³", "0"],
    ],
    [
        ["Center", "TCAD type", "Energy level", "σe (cm²)", "σh (cm²)", "Trap density", "Source"],
        ["Z1/2", "Acceptor", "Ec - 0.67 eV", "2 × 10⁻¹⁴", "1 × 10⁻¹⁵", "Nt", "Gaggl et al.; Capan"],
        ["EH6/7", "Donor", "Ec - 1.55 eV", "2 × 10⁻¹⁴", "1 × 10⁻¹⁵", "Nt", "Kleppinger et al.; Gaggl et al."],
    ],
    [
        ["Energy", "ESTAR RCSDA (μm)", "Geant4 z50 (μm)", "Geant4 z90 (μm)", "Edep / Ein"],
        ["20 keV", "3.44", "0.98", "1.95", "0.918"],
        ["49 keV", "16.35", "4.93", "9.56", "0.924"],
        ["100 keV", "55.36", "16.87", "32.61", "0.929"],
        ["156.5 keV", "115.94", "35.09", "67.23", "0.934"],
    ],
    [
        ["Nt (cm⁻³)", "20 keV", "49 keV", "100 keV", "156.5 keV", "¹⁴C"],
        ["0", "110 / 74.24", "130 / 97.74", "120 / 99.41", "150 / 99.62", "150 / 97.80"],
        ["10¹²", "110 / 74.19", "130 / 97.67", "120 / 99.35", "150 / 99.56", "100 / 97.72"],
        ["10¹³", "8 / 73.88", "20 / 97.48", "60 / 97.39", "130 / 72.37", "60 / 95.09"],
        ["2.5 × 10¹³", "8 / 73.82", "20 / 97.21", "40 / 80.54", "180 / 52.03", "40 / 87.14"],
        ["5 × 10¹³", "5 / 73.78", "20 / 96.59", "30 / 66.87", "180 / 37.69", "30 / 81.72"],
    ],
]


def table_font_size(index: int) -> float:
    if index == 2:
        return 9.2
    if index == 3:
        return 8.8
    if index == 5:
        return 8.7
    return 9.5


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9.5, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    clear_paragraph_indent(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold)


def rebuild_table_content(table, index: int) -> None:
    data = TABLE_DATA[index - 1]
    rows_needed = len(data)
    cols_needed = len(data[0])

    while len(table.rows) < rows_needed:
        table.add_row()
    while len(table.rows) > rows_needed:
        table._tbl.remove(table.rows[-1]._tr)

    # Pandoc already creates the correct column count for these tables. If it
    # ever changes, fail loudly rather than silently producing a malformed table.
    if len(table.rows[0].cells) != cols_needed:
        raise RuntimeError(f"Table {index} has {len(table.rows[0].cells)} columns, expected {cols_needed}")

    for r, row_data in enumerate(data):
        for c, value in enumerate(row_data):
            set_cell_text(table.rows[r].cells[c], value, bold=(r == 0), size=table_font_size(index))


def format_table(table, index: int) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    rebuild_table_content(table, index)
    set_tbl_borders(table)
    set_table_row_rules(table)
    if index in (2, 5):
        set_cell_margins(table, top=55, start=110, bottom=55, end=110)
    elif index == 3:
        set_cell_margins(table, top=60, start=110, bottom=60, end=110)
    else:
        set_cell_margins(table)

    widths_by_table = {
        1: [2.45, 3.75],
        2: [1.25, 1.65, 1.65, 1.35],
        3: [0.80, 0.95, 0.95, 0.65, 0.65, 0.90, 1.52],
        4: [1.05, 1.45, 1.35, 1.35, 1.05],
        5: [1.10, 1.04, 1.04, 1.04, 1.04, 1.04],
    }
    widths = widths_by_table.get(index)
    if widths:
        set_table_geometry(table, widths)

    for row_i, row in enumerate(table.rows):
        prevent_row_split(row)
        for col_i, cell in enumerate(row.cells):
            clear_cell_shading(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                clear_paragraph_indent(paragraph)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size_pt=table_font_size(index), bold=(row_i == 0))


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def paragraph_has_math(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//*[local-name()='oMath' or local-name()='oMathPara']"))


def paragraph_text(paragraph) -> str:
    return re.sub(r"\s+", " ", paragraph.text.strip())


def is_figure_caption(text: str) -> bool:
    return bool(re.match(r"^Fig\.\s*\d+\.", text))


def is_table_caption(text: str) -> bool:
    return bool(re.match(r"^Table\s+\d+\.", text))


def format_table_caption(paragraph) -> None:
    text = paragraph_text(paragraph)
    match = re.match(r"^(Table\s+[0-9]+)\.\s*(.+)$", text)
    if not match:
        return

    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    clear_paragraph_indent(paragraph)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)

    number_run = paragraph.add_run(match.group(1))
    set_run_font(number_run, size_pt=10.5, bold=True)
    number_run.add_break()
    title_run = paragraph.add_run(match.group(2))
    set_run_font(title_run, size_pt=10.5, bold=False)


def apply_paragraph_styles(doc: Document) -> None:
    in_references = False
    in_abstract = False

    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)

        if paragraph_has_drawing(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clear_paragraph_indent(paragraph)
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.keep_with_next = True
            continue

        if text == "Abstract":
            in_abstract = True
        if text == "References":
            in_references = True

        if is_figure_caption(text):
            paragraph.style = doc.styles["Caption"]
            clear_paragraph_indent(paragraph)
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.line_spacing = 1.03
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif is_table_caption(text):
            paragraph.style = doc.styles["Caption"]
            format_table_caption(paragraph)
        elif in_references and re.match(r"^\[[0-9]+\]", text):
            clear_paragraph_indent(paragraph)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.03
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_run_font(run, size_pt=10.5)
        elif text.startswith("Keywords:"):
            in_abstract = False
            clear_paragraph_indent(paragraph)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif paragraph.style.name in ("Normal", "Body Text", "First Paragraph", "FirstParagraph"):
            if paragraph_has_math(paragraph) and len(text) < 120:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                clear_paragraph_indent(paragraph)
            elif in_abstract:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                clear_paragraph_indent(paragraph)
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.first_line_indent = Inches(0.28)
            paragraph.paragraph_format.space_after = Pt(3.5)
            paragraph.paragraph_format.line_spacing = 1.08

        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            clear_paragraph_indent(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if paragraph.style.name == "Title":
            remove_paragraph_border(paragraph)
            clear_paragraph_indent(paragraph)

        for run in paragraph.runs:
            if paragraph.style.name == "Title":
                set_run_font(run, size_pt=16, bold=True, italic=False)
            elif paragraph.style.name == "Heading 1":
                set_run_font(run, size_pt=14, bold=True, italic=False)
            elif paragraph.style.name == "Heading 2":
                set_run_font(run, size_pt=13, bold=True, italic=False)
            elif paragraph.style.name == "Heading 3":
                set_run_font(run, size_pt=12, bold=True, italic=False)
            elif paragraph.style.name == "Caption" and not paragraph_text(paragraph).startswith("Table "):
                set_run_font(run, size_pt=10.5, bold=False)
            elif not in_references:
                set_run_font(run, size_pt=12)


def resize_figures(doc: Document) -> None:
    figure_specs = [
        ("fig1_sic_pin_structure.png", 3.20),
        ("fig2_geant4_tcad_workflow.png", 5.15),
        ("fig3_tcad_generation_distribution.png", 3.95),
        ("fig4_c14_spectrum.png", 3.65),
        ("fig5_dedx_profiles.png", 3.80),
        ("fig6_cv_1overc2_baseline.png", 4.20),
        ("fig7_c14_cce_vs_thickness_by_Nt.png", 4.65),
        ("fig8_c14_cce_design_map.png", 4.15),
        ("fig9_optimal_thickness_vs_Nt.png", 3.50),
        ("fig10_design_bias_matrix.png", 2.55),
    ]
    figure_dir = BASE_DIR.parent / "figures"
    for shape, (filename, width_in) in zip(doc.inline_shapes, figure_specs):
        image_path = figure_dir / filename
        with Image.open(image_path) as image:
            pixel_width, pixel_height = image.size
        height_in = width_in * pixel_height / pixel_width
        shape.width = Inches(width_in)
        shape.height = Inches(height_in)


def insert_page_break_before(paragraph) -> None:
    new_p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    new_p.append(r)
    paragraph._p.addprevious(new_p)


def insert_page_break_after(paragraph) -> None:
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def insert_page_break_at_start(paragraph) -> None:
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, r)


def group_fig10_image_caption(doc: Document) -> None:
    fig_paragraph = None
    caption_paragraph = None
    drawing_index = 0
    for paragraph in doc.paragraphs:
        if paragraph_has_drawing(paragraph):
            drawing_index += 1
            if drawing_index == 10:
                fig_paragraph = paragraph
        elif paragraph_text(paragraph).startswith("Fig.10."):
            caption_paragraph = paragraph
            break

    if fig_paragraph is None or caption_paragraph is None:
        return

    image_path = BASE_DIR.parent / "figures" / "fig10_design_bias_matrix.png"
    if not image_path.exists():
        return

    caption_text = paragraph_text(caption_paragraph)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    clear_table_borders(table)
    set_cell_margins(table, top=0, start=0, bottom=0, end=0)
    set_table_geometry(table, [5.90])

    image_cell = table.rows[0].cells[0]
    image_p = image_cell.paragraphs[0]
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_paragraph_indent(image_p)
    image_p.paragraph_format.space_before = Pt(4)
    image_p.paragraph_format.space_after = Pt(2)
    with Image.open(image_path) as image:
        pixel_width, pixel_height = image.size
    width_in = 2.55
    height_in = width_in * pixel_height / pixel_width
    image_p.add_run().add_picture(str(image_path), width=Inches(width_in), height=Inches(height_in))

    caption_p = image_cell.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    clear_paragraph_indent(caption_p)
    caption_p.paragraph_format.space_before = Pt(2)
    caption_p.paragraph_format.space_after = Pt(5)
    caption_p.paragraph_format.line_spacing = 1.03
    caption_run = caption_p.add_run(caption_text)
    set_run_font(caption_run, size_pt=10.5)

    fig_paragraph._p.addprevious(table._tbl)
    fig_paragraph._element.getparent().remove(fig_paragraph._element)
    caption_paragraph._element.getparent().remove(caption_paragraph._element)


def add_targeted_page_breaks(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph_text(paragraph).startswith("Table 1") or paragraph_text(paragraph).startswith("Table 2"):
            paragraph.paragraph_format.page_break_before = True


def postprocess_docx() -> None:
    doc = Document(PANDOC_DOCX)

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.20)
    section.bottom_margin = Cm(2.20)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    styles = doc.styles
    configure_style(styles["Normal"], size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=3.5, line=1.08, first_line=0.28)
    configure_style(styles["Title"], size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12, line=1.08)
    remove_style_border(styles["Title"])
    configure_style(styles["Heading 1"], size=14, bold=True, before=12, after=5, line=1.08, keep_next=True)
    configure_style(styles["Heading 2"], size=13, bold=True, before=10, after=4, line=1.08, keep_next=True)
    configure_style(styles["Heading 3"], size=12, bold=True, italic=False, before=7, after=3, line=1.08, keep_next=True)
    configure_style(styles["Caption"], size=10.5, before=2, after=5, line=1.03)
    if "Body Text" in styles:
        configure_style(styles["Body Text"], size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=3.5, line=1.08, first_line=0.28)
    if "First Paragraph" in styles:
        configure_style(styles["First Paragraph"], size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=3.5, line=1.08, first_line=0.28)
    if "FirstParagraph" in styles:
        configure_style(styles["FirstParagraph"], size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=3.5, line=1.08, first_line=0.28)

    apply_paragraph_styles(doc)
    resize_figures(doc)
    for idx, table in enumerate(doc.tables, start=1):
        format_table(table, idx)
    group_fig10_image_caption(doc)
    add_targeted_page_breaks(doc)

    doc.save(CANDIDATE_DOCX)
    Document(CANDIDATE_DOCX)
    if OUTPUT_DOCX.exists():
        OUTPUT_DOCX.unlink()
    CANDIDATE_DOCX.replace(OUTPUT_DOCX)


def main() -> None:
    BASE_DIR.mkdir(parents=True, exist_ok=True)
    prepare_docx_markdown()
    build_reference_docx()
    run_pandoc()
    postprocess_docx()
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
